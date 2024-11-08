import os
import pandas as pd
import docx
from docx.shared import Pt, Cm, Mm
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from Image_Behind import add_float_picture
from docx.table import Table

# Custom class for caching cells in a table for performance improvement
class CachedTable(Table):
    def __init__(self, tbl, parent):
        super(Table, self).__init__(parent)
        self._element = self._tbl = tbl
        self._cached_cells = None

    # Property to retrieve or cache table cells
    @property
    def _cells(self):
        if self._cached_cells is None:
            self._cached_cells = super(CachedTable, self)._cells
        return self._cached_cells

    # Static method to convert a table to a CachedTable
    @staticmethod
    def transform(table):
        cached_table = CachedTable(table._tbl, table._parent)
        return cached_table

start_time = time.time()  # Start timing the script

# Initialize a new Word document
doc = docx.Document()

# Set page margins and dimensions for all sections
margin = 2.54
sections = doc.sections
for section in sections:
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

# Add header image to document header
header = section.header
paragraph = header.paragraphs[0]
add_float_picture(paragraph, 'Header.png', width=Cm(3.6), pos_x=Cm(8.86), pos_y=Cm(0.09))

# Add footer text and icons
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.text = "\t     www.centroveda.pt              +351 239 701 704             centroveda@centroveda.pt"
paragraph.style.font.size = Pt(12)
add_float_picture(paragraph, 'Icon web.png', width=Cm(1.06), pos_x=Cm(3.1), pos_y=Cm(27.69))
add_float_picture(paragraph, 'Icon phone.png', width=Cm(1.06), pos_x=Cm(7.86), pos_y=Cm(27.69))
add_float_picture(paragraph, 'Icon mail.png', width=Cm(1.06), pos_x=Cm(12.19), pos_y=Cm(27.69))

# Function to set column widths in the table
def set_col_widths(table):
    widths = (2.7, 2.75, 0.55, 2.75, 0.55, 2.2, 2.5, 2.8)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Cm(width)

# Load data from Excel file
dft = pd.read_excel(r'RETENTORES LISTAGEM GERAL - truncated.xlsx')
df = dft

# Remove unwanted column names from data
df.columns = df.columns.str.replace("Unnamed: 0", "")
df.columns = df.columns.str.replace("Unnamed: 2", "")
df.columns = df.columns.str.replace("Unnamed: 4", "")

# Split data into chunks of 46 rows for easier processing
dfs = [df.iloc[i:i+46] for i in range(0, len(df), 46)]
counter = 0

# Iterate over each data chunk to create a table in the document
for df in dfs:
    counter += 1

    # Initialize a new CachedTable with the data's number of rows and columns
    t = CachedTable.transform(doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1]))
    table_cells = t._cells
    set_col_widths(t)

    # Add header row with column names
    for j in range(df.shape[1]):
        t.cell(0, j).text = df.columns[j]
        paragraph = t.cell(0, j).paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add data rows to the table
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = df.iat[i, j]
            t.cell(i + 1, j).text = str(cell)

    # Apply specific formatting based on cell text
    for i in range(df.shape[0]):
        for j, each_cell in enumerate(t.rows[i + 1].cells):
            if each_cell.text == "NBR":
                each_cell.text = ""
                paragraph = each_cell.paragraphs[0]
                run = paragraph.add_run("NBR")
                run.font.bold = True
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif each_cell.text == "VITON":
                each_cell.text = ""
                paragraph = each_cell.paragraphs[0]
                run = paragraph.add_run("VITON")
                run.font.color.rgb = RGBColor(131, 60, 12)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif each_cell.text == "SILICONE":
                each_cell.text = ""
                paragraph = each_cell.paragraphs[0]
                run = paragraph.add_run("SILICONE")
                run.font.color.rgb = RGBColor(206, 33, 20)
                run.font.bold = True
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                val = each_cell.text
                each_cell.text = ""
                paragraph = each_cell.paragraphs[0]
                run = paragraph.add_run(val)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set default font and style for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    t.style = 'Light List Accent 1'

    # Add a background image below each table
    p = doc.add_paragraph()
    add_float_picture(p, 'Fundo.png', width=Cm(13.6), pos_x=Cm(3.7), pos_y=Cm(5.85))

    # Add a new section if there are more tables to process
    if counter != len(dfs):
        doc.add_section()

# Save the final document
doc.save('Listagem Retentores W.docx')

# Print the time taken to execute the script
print("--- %s seconds ---" % (time.time() - start_time))
